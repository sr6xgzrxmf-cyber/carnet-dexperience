from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('/Users/laurent/carnet-dexperience/outputs/01a05798-5711-7ac1-93aa-232f24b67032/questionnaire-interview-portfolio.docx')
NAVY = '252854'; TEAL = '11BFA5'; PALE = 'E9F7F3'; LIGHT = 'F3F3EF'; MID = '6B6E78'; WHITE = 'FFFFFF'; INK = '20222A'; AMBER = 'F4E6C7'

tiles = [
('01', 'Faire vivre un projet en construisant son écosystème', 'Le jour où j’ai arrêté de chercher des mécènes', [
 'Revenez au tout début : quel était le projet, quelle était votre mission officielle et quel résultat cherchiez-vous alors ?',
 'Quel obstacle concret vous a obligé à dépasser la seule recherche de financement ?',
 'Qu’avez-vous personnellement mis en place pour structurer les contacts, les relances et les relais ?',
 'Quel moment vous a fait comprendre que vous construisiez un écosystème plutôt qu’une liste de mécènes ?',
 'Quels résultats pouvez-vous attribuer à cette démarche : relais, presse, réservations, audience, partenariats ?',
 'Quels supports ou traces pouvez-vous montrer : tableau, mails, QR codes, dossier, interview, publications ?',
 'Si cette expérience devait prouver une seule compétence professionnelle, laquelle choisiriez-vous et pourquoi ?']),
('02', 'Quitter son argumentaire pour écouter les contraintes de l’autre', 'Le moment où j’ai compris que mon discours ne fonctionnait pas', [
 'Dans quel contexte réel cette conversation s’est-elle déroulée, avec quels interlocuteurs et quel enjeu ?',
 'Quels signaux vous ont fait comprendre que votre discours ne rencontrait pas le besoin de l’autre ?',
 'Qu’avez-vous ressenti et quel risque preniez-vous en abandonnant votre argumentaire préparé ?',
 'Quelle question ou reformulation a réellement changé la conversation ?',
 'Qu’avez-vous appris sur les contraintes ou critères de décision de votre interlocuteur ?',
 'Qu’est-ce qui a changé ensuite dans votre manière de préparer et conduire ces échanges ?',
 'Quelle preuve ou quel exemple ultérieur montre que vous avez durablement intégré cet apprentissage ?']),
('03', 'Transformer un refus financier en relais concret', 'Quand un partenaire dit non…', [
 'Que demandiez-vous exactement et quelle forme a pris le refus ?',
 'Quel détail dans la réponse vous a permis de voir qu’une autre possibilité restait ouverte ?',
 'Quelle alternative précise avez-vous proposée, et pourquoi celle-ci plutôt qu’une autre ?',
 'Quels contenus ou moyens avez-vous fournis pour rendre le relais facile à activer ?',
 'Que s’est-il réellement passé après votre réponse ?',
 'Qu’avez-vous changé dans votre façon de lire et classer les réponses reçues ?',
 'Quelles traces permettraient de documenter ce passage du refus au relais ?']),
('04', 'Adapter le langage sans déformer le projet', 'Construire un partenariat sans devenir commercial', [
 'Quel était le projet et quelle adaptation vous était suggérée ?',
 'À quel moment avez-vous senti que la traduction risquait de devenir une déformation ?',
 'Quelles valeurs ou dimensions du projet étaient non négociables pour vous ?',
 'Comment avez-vous posé cette limite sans fermer la relation ?',
 'Quelle a été la réaction des interlocuteurs et quelle décision a suivi ?',
 'Qu’est-ce que cette situation dit de votre manière de négocier ou de construire un partenariat ?',
 'Quelle preuve montre que le projet est resté reconnaissable tout en devenant compréhensible ?']),
('05', 'Réduire une vente immédiate pour sécuriser l’adoption', 'La vente que je n’ai pas voulu conclure', [
 'Quels éléments de ce récit recomposé proviennent directement de situations professionnelles réelles ?',
 'Quel besoin le client formulait-il et quel problème plus profond avez-vous diagnostiqué ?',
 'Quelles questions vous ont permis de déplacer la discussion du produit vers l’usage ?',
 'Pourquoi avez-vous recommandé un déploiement plus limité, malgré l’opportunité commerciale ?',
 'Comment avez-vous accompagné le test ou les premiers utilisateurs ?',
 'Quels effets avez-vous pu observer sur l’adoption, la décision ou la relation client ?',
 'Quelles preuves sont utilisables sans révéler d’informations confidentielles ?']),
('06', 'Faire coopérer sans autorité hiérarchique', 'Déménager une équipe sans en être le manager', [
 'Qui composait cette équipe informelle et quelles différences de capacités ou de tempérament deviez-vous intégrer ?',
 'Quel était le risque principal au démarrage : désorganisation, fatigue, tension, retard ou autre ?',
 'Comment avez-vous rendu l’objectif commun suffisamment clair ?',
 'Racontez un moment précis où vous avez redistribué une tâche, rassuré quelqu’un ou régulé une tension.',
 'Qu’avez-vous volontairement laissé faire sans intervenir ?',
 'Quels signes vous ont montré que le collectif fonctionnait et que l’entraide avait émergé ?',
 'À quelle situation professionnelle cette expérience fait-elle écho dans votre parcours ?']),
('07', 'Transformer un site en système éditorial durable', 'Apprendre à bâtir un média, pas un site', [
 'Quel était l’état initial de Carnet d’expérience et qu’est-ce qui vous paraissait encore insuffisant ?',
 'Quelle décision a marqué le passage du simple site au média structuré ?',
 'Qu’avez-vous appris ou construit vous-même sur les plans technique, éditorial et organisationnel ?',
 'Comment avez-vous organisé les articles, séries, archives, rythmes et liens entre contenus ?',
 'Quelles difficultés ou erreurs ont entraîné une modification durable de votre méthode ?',
 'Quels résultats sont aujourd’hui visibles ou mesurables sur le site ?',
 'Quelles captures, versions, fichiers ou statistiques peuvent servir de preuves ?']),
('08', 'Protéger la cohérence d’un média par des renoncements explicites', 'Les choix que j’ai faits et ceux que j’ai refusés', [
 'Quelles fonctionnalités, stratégies ou formats avez-vous réellement envisagés ?',
 'Quels critères vous ont permis de décider ce qui appartenait ou non au projet ?',
 'Quel refus a été le plus difficile, et qu’auriez-vous peut-être gagné en disant oui ?',
 'Qu’avez-vous protégé grâce à ces renoncements : ton, confiance, accessibilité, rythme, cohérence ?',
 'Comment avez-vous traduit ces choix dans la structure concrète du site ?',
 'Avez-vous constaté un effet sur les lecteurs, votre écriture ou votre manière de travailler ?',
 'Quelle décision illustre le mieux votre capacité à arbitrer sans perdre l’intention initiale ?']),
('09', 'Transformer une fatigue récurrente en méthode de travail', 'De la fatigue à la méthode', [
 'Quelle tâche ou situation se répétait au point de devenir fatigante ?',
 'Comment avez-vous compris que le problème venait du système plutôt que d’un manque d’effort ?',
 'Qu’avez-vous observé, noté ou mesuré avant de concevoir une méthode ?',
 'Décrivez la méthode, la routine ou le support que vous avez créé.',
 'Qui l’a utilisé, dans quelles conditions et avec quelles adaptations ?',
 'Quels gains avez-vous observés : temps, qualité, sérénité, autonomie ou fiabilité ?',
 'Le support existe-t-il encore sous une forme montrable ou reconstituable ?']),
('10', 'Passer d’un inventaire d’expériences à un système de décision', 'Quand un parcours devient pilotable', [
 'À quel moment votre parcours vous est-il apparu riche mais difficile à lire ?',
 'Quelles informations avez-vous réunies pour commencer à le cartographier ?',
 'Quels regroupements, axes ou critères vous ont permis de faire émerger une cohérence ?',
 'Qu’avez-vous écarté, hiérarchisé ou reformulé pendant ce travail ?',
 'Quelles décisions professionnelles sont devenues plus faciles grâce à cette nouvelle lecture ?',
 'Quels documents ou outils matérialisent cette démarche ?',
 'Comment cette expérience pourrait-elle devenir une méthode utile pour d’autres personnes ?']),
('11', 'Rendre visible une contribution transversale', 'Faire reconnaître un rôle transversal sans organigramme', [
 'Dans quelle organisation avez-vous tenu un rôle utile mais insuffisamment nommé ?',
 'Quelles situations déclenchaient naturellement votre intervention ?',
 'Que faisiez-vous concrètement pour clarifier, relier, débloquer ou fluidifier ?',
 'Racontez une situation où votre intervention a changé la compréhension ou la coordination.',
 'Quels effets revenaient régulièrement lorsque vous interveniez ?',
 'Comment ce rôle était-il perçu par les équipes, les managers ou les partenaires ?',
 'Quelles preuves indirectes existent : sollicitations, messages, responsabilités confiées, résultats ?']),
('12', 'Traiter une erreur répétée comme un problème de système', 'Améliorer un processus sans conflit', [
 'Avez-vous vécu une situation où les mêmes erreurs revenaient entre personnes ou équipes ?',
 'Quels reproches ou tensions masquaient alors le véritable problème de processus ?',
 'Comment avez-vous recueilli les faits sans chercher un responsable ?',
 'Avec qui avez-vous travaillé pour comprendre les besoins en amont et en aval ?',
 'Quel outil, standard, contrôle ou check-list avez-vous conçu ou proposé ?',
 'Quels résultats ont été observés après sa mise en place ?',
 'Quelle preuve permettrait de montrer l’amélioration sans exposer de données sensibles ?']),
('13', 'Sortir d’un débat en organisant un test réversible', 'Le prototype 30 jours', [
 'Quelle idée ou décision réelle aurait mérité un test limité à trente jours ?',
 'Pourquoi la discussion seule ne permettait-elle plus d’avancer ?',
 'Quel périmètre minimal auriez-vous choisi pour rendre le test sûr et utile ?',
 'Quelles règles et limites auraient protégé les personnes concernées ?',
 'Quels deux ou trois indicateurs auraient permis de décider ?',
 'Qui aurait assuré le suivi et à quel rythme ?',
 'Avez-vous déjà utilisé une démarche comparable, même sans l’appeler “prototype 30 jours” ?']),
('14', 'Transformer les savoirs dispersés en mémoire de travail', 'Création collaborative d’un site de formation interne', [
 'Ce projet correspond-il à une réalisation vécue, une proposition ou une réflexion théorique ?',
 'Quels savoirs risquaient de rester informels ou de disparaître ?',
 'Comment les personnes du terrain ont-elles été associées à la sélection des contenus ?',
 'Quel était votre rôle exact : pilotage, cadrage, édition, production, coordination ?',
 'Comment avez-vous transformé l’expertise brute en contenu compréhensible et utilisable ?',
 'Qu’est-ce qui a été effectivement produit et utilisé ?',
 'Quels retours, contenus ou usages peuvent servir de preuves ?']),
('15', 'Concevoir la transmission pour rendre autonome', 'Une formation réussie commence quand le formateur devient inutile', [
 'Choisissez une formation réelle dont l’objectif était l’autonomie, pas seulement la compréhension.',
 'Que devaient savoir faire les participants après la formation ?',
 'Comment avez-vous adapté le contenu aux expériences et difficultés réelles du groupe ?',
 'Quelles activités ont permis aux participants d’agir sans dépendre du formateur ?',
 'Comment avez-vous vérifié le transfert vers le terrain ?',
 'Quels changements ou résultats avez-vous observés après la formation ?',
 'Quelles preuves existent : support, évaluation, témoignage, production ou comportement observable ?']),
]

def set_font(run, name='Aptos', size=10.5, color=INK, bold=False, italic=False):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),color)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_table_geometry(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            cell.width=Inches(widths[i]/1440); set_cell_margins(cell)
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            tcW.set(qn('w:w'),str(widths[i])); tcW.set(qn('w:type'),'dxa')

def no_borders(table):
    tblPr=table._tbl.tblPr; borders=tblPr.find(qn('w:tblBorders'))
    if borders is None: borders=OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'nil'); borders.append(el)

def bottom_border(paragraph, color='DADBE1', size='8'):
    pPr=paragraph._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),size); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),color); pBdr.append(bottom); pPr.append(pBdr)

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=paragraph.add_run('Page '); set_font(r,size=8.5,color=MID)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)

def add_question(doc, number, question):
    t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); no_borders(t)
    c=t.cell(0,0); shade(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    r=p.add_run(f'{number}.  {question}'); set_font(r,size=9.6,bold=True)
    a=doc.add_paragraph(); a.paragraph_format.space_before=Pt(2); a.paragraph_format.space_after=Pt(5)
    rr=a.add_run('________________________________________________________________________________'); set_font(rr,size=8,color='B7B8BC')
    a2=doc.add_paragraph(); a2.paragraph_format.space_before=Pt(0); a2.paragraph_format.space_after=Pt(5)
    rr=a2.add_run('________________________________________________________________________________'); set_font(rr,size=8,color='B7B8BC')
    a3=doc.add_paragraph(); a3.paragraph_format.space_before=Pt(0); a3.paragraph_format.space_after=Pt(6)
    rr=a3.add_run('________________________________________________________________________________'); set_font(rr,size=8,color='B7B8BC')

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.2677); sec.page_height=Inches(11.6929); sec.top_margin=Inches(.62); sec.bottom_margin=Inches(.58); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72); sec.header_distance=Inches(.32); sec.footer_distance=Inches(.32)

normal=doc.styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
for name,size,before,after,color in [('Heading 1',16,18,10,NAVY),('Heading 2',13,14,7,NAVY),('Heading 3',11,10,5,'1F4D78')]:
    st=doc.styles[name]; st.font.name='Aptos'; st._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.text='CARNET D’EXPÉRIENCE   ·   INTERVIEW PORTFOLIO'; set_font(header.runs[0],size=8.2,color=MID,bold=True); bottom_border(header)
add_page_number(sec.footer.paragraphs[0])

# Cover - editorial cover override for a practical interview workbook.
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('MOSAÏQUE NARRATIVE'); set_font(r,size=10,color=TEAL,bold=True); p.paragraph_format.space_after=Pt(18)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Carnet d’interview\ndu portfolio'); set_font(r,name='Georgia',size=30,color=NAVY,bold=True); p.paragraph_format.space_after=Pt(12)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('15 récits à faire émerger à partir des articles de Carnet d’expérience'); set_font(r,size=13,color=MID); p.paragraph_format.space_after=Pt(46)
t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); no_borders(t); c=t.cell(0,0); shade(c,PALE)
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); r=p.add_run('Contexte  →  problème  →  intervention  →  résultat  →  compétences  →  preuves'); set_font(r,size=11,color=NAVY,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50); r=p.add_run('Laurent Guyonnet  ·  Août 2026'); set_font(r,size=10,color=MID)
doc.add_page_break()

# Method page.
p=doc.add_paragraph(); r=p.add_run('Comment utiliser ce carnet'); set_font(r,name='Georgia',size=25,color=NAVY,bold=True); p.paragraph_format.space_after=Pt(8)
p=doc.add_paragraph('Ce document n’est pas un exercice de rédaction. Il sert à retrouver les faits, les gestes, les décisions et les effets qui permettront ensuite d’écrire chaque tuile avec précision.'); p.paragraph_format.space_after=Pt(14)
for title,text in [
('1. Parler avant de formuler','Répondez comme vous raconteriez la situation à quelqu’un. Les hésitations et les détails concrets sont souvent plus utiles qu’une réponse déjà “professionnelle”.'),
('2. Distinguer le collectif de votre contribution','Le projet peut être collectif. Il faut néanmoins identifier ce que vous avez personnellement observé, décidé, proposé, construit ou rendu possible.'),
('3. Chercher les effets observables','Un résultat n’est pas obligatoirement un chiffre. Il peut s’agir d’une décision prise, d’une tension réduite, d’un usage adopté, d’un relais obtenu ou d’une autonomie nouvelle.'),
('4. Ne pas inventer la preuve','Un article peut exprimer une conviction sans démontrer une expérience. Si la situation n’a pas été vécue, notez-le : elle pourra devenir une tuile “méthode” ou rester un article lié.'),
('5. Une fiche à la fois','Chaque entretien peut porter sur une seule tuile. Vingt à trente minutes suffisent généralement pour produire la matière d’un premier récit.')]:
    t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); no_borders(t); c=t.cell(0,0); shade(c,LIGHT)
    pp=c.paragraphs[0]; rr=pp.add_run(title); set_font(rr,size=11,color=NAVY,bold=True); pp.add_run('\n'); rr=pp.add_run(text); set_font(rr,size=10,color=INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
p=doc.add_paragraph(); r=p.add_run('À garder sous les yeux pendant chaque entretien'); set_font(r,size=11,color=TEAL,bold=True)
p=doc.add_paragraph('Qui ?  ·  Où ?  ·  Quand ?  ·  Qu’est-ce qui bloquait ?  ·  Qu’avez-vous fait ?  ·  Qu’est-ce qui a changé ?  ·  Comment le prouver ?'); p.paragraph_format.space_after=Pt(0)
doc.add_page_break()

for idx,(num,title,source,questions) in enumerate(tiles):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run(f'TUILE {num}   ·   QUESTIONNAIRE GUIDÉ'); set_font(r,size=9,color=TEAL,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); r=p.add_run(title); set_font(r,name='Georgia',size=21,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(9); r=p.add_run(f'Article source : {source}'); set_font(r,size=9.2,color=MID,italic=True)
    for qi,q in enumerate(questions,1): add_question(doc,qi,q)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(0); r=p.add_run('Phrase à retenir — si vous ne deviez conserver qu’une seule idée de cette histoire :'); set_font(r,size=9.5,color=TEAL,bold=True)
    a=doc.add_paragraph('________________________________________________________________________________'); a.paragraph_format.space_after=Pt(0); set_font(a.runs[0],size=8,color='B7B8BC')
    if idx < len(tiles)-1: doc.add_page_break()

doc.core_properties.title='Carnet d’interview du portfolio - Mosaïque narrative'
doc.core_properties.subject='Questionnaire guidé pour rédiger quinze tuiles de portfolio'
doc.core_properties.author='Laurent Guyonnet'
doc.core_properties.keywords='portfolio, mosaïque narrative, compétences, expériences, preuves'
doc.save(OUT)
print(OUT)
